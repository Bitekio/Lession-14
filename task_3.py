"""
hw
"""

from docx import Document
from docx.shared import Pt, RGBColor

if __name__ == "__main__":
    doc = Document()
    doc.add_paragraph("Hello Python")
    doc.save("hello_python.docx")

    read_doc = Document("hello_python.docx")
    bold_text = [run.text for para in read_doc.paragraphs for run in para.runs if run.bold]
    print("Bold Text in Python String:", " ".join(bold_text))

    new_doc = Document()
    new_paragraph = new_doc.add_paragraph("New Paragraph with Custom Font")
    run = new_paragraph.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 255)

    new_doc.save("custom_paragraph.docx")
